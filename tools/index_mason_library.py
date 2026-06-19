"""Build a local Mason/梅森 research-library index.

The script reads local PDFs, Word files, Excel sheets, and RTF notes, then
creates a lightweight JSON/Markdown index for the workbench. It stores only
metadata, short previews, hashes, and theme tags. It does not copy source files.
"""

from __future__ import annotations

import argparse
import hashlib
import json
import re
from collections import Counter
from datetime import datetime
from pathlib import Path
from typing import Any


THEMES: dict[str, list[str]] = {
    "AI算力/国产算力": [
        "算力",
        "国产算力",
        "GPU",
        "英伟达",
        "NVIDIA",
        "HBM",
        "服务器",
        "液冷",
        "电源",
        "中芯国际",
        "寒武纪",
        "华为",
        "GB300",
        "超算",
    ],
    "CPO/光通信": [
        "CPO",
        "光模块",
        "光通信",
        "光互联",
        "OCS",
        "光交换",
        "光纤",
        "光电",
        "长飞",
        "亨通",
        "中际",
        "新易盛",
        "天孚",
    ],
    "PCB/ABF/载板": [
        "PCB",
        "ABF",
        "载板",
        "兴森",
        "沪电",
        "生益",
        "深南",
        "玻璃基板",
    ],
    "半导体/材料/设备": [
        "半导体",
        "晶圆",
        "DSP",
        "功率",
        "士兰微",
        "扬杰",
        "存储",
        "材料",
        "设备",
        "国产化",
        "芯片",
    ],
    "AI应用/Agent": [
        "AI应用",
        "Agent",
        "Manus",
        "DeepSeek",
        "大模型",
        "应用",
        "医疗",
        "软件",
    ],
    "机器人/无人驾驶": [
        "机器人",
        "宇树",
        "灵巧手",
        "无人驾驶",
        "智能驾驶",
        "禾赛",
    ],
    "小金属/有色/稀土": [
        "小金属",
        "有色",
        "稀土",
        "铋",
        "锑",
        "钨",
        "锂",
        "锆",
        "钽",
        "黄金",
        "铜",
        "铝",
    ],
    "能源/化工/油价": [
        "能源",
        "化工",
        "油价",
        "原油",
        "缺电",
        "电力",
        "核聚变",
        "固态电池",
        "锂电池",
    ],
    "宏观/策略/风格": [
        "宏观",
        "策略",
        "风格",
        "监管",
        "关税",
        "冲突",
        "美伊",
        "以伊",
        "市场",
        "行情",
        "展望",
        "抄底",
        "黄金坑",
        "攻与守",
    ],
    "消费/其他": ["消费", "泡泡", "黄金珠宝", "海洋经济", "商业航天", "墨脱"],
}

SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".xlsx", ".xls", ".rtf"}


def clean_text(text: str | None) -> str:
    if not text:
        return ""
    return re.sub(r"\s+", " ", text).strip()


def file_hash(path: Path) -> str:
    digest = hashlib.md5()
    with path.open("rb") as handle:
        for chunk in iter(lambda: handle.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest()


def infer_themes(name: str, sample_text: str = "") -> list[str]:
    haystack = f"{name} {sample_text[:3000]}".lower()
    scores: dict[str, int] = {}
    for theme, keywords in THEMES.items():
        score = sum(haystack.count(keyword.lower()) for keyword in keywords)
        if score:
            scores[theme] = score
    if not scores:
        return ["未分类"]
    return [theme for theme, _ in sorted(scores.items(), key=lambda item: item[1], reverse=True)[:3]]


def extract_pdf(path: Path, max_pages: int) -> dict[str, Any]:
    result: dict[str, Any] = {"pages": None, "sample_text": "", "error": None}
    try:
        import pdfplumber  # type: ignore

        with pdfplumber.open(str(path)) as pdf:
            result["pages"] = len(pdf.pages)
            texts = [(page.extract_text() or "") for page in pdf.pages[:max_pages]]
            result["sample_text"] = clean_text("\n".join(texts))[:6000]
    except Exception as exc:  # pragma: no cover - depends on local PDFs
        result["error"] = repr(exc)
    return result


def extract_docx(path: Path) -> dict[str, Any]:
    result: dict[str, Any] = {"paragraphs": None, "sample_text": "", "error": None}
    try:
        import docx  # type: ignore

        document = docx.Document(str(path))
        paragraphs = [p.text for p in document.paragraphs if p.text.strip()]
        for table in document.tables[:3]:
            for row in table.rows[:20]:
                paragraphs.append(" | ".join(cell.text.strip() for cell in row.cells))
        result["paragraphs"] = len(paragraphs)
        result["sample_text"] = clean_text("\n".join(paragraphs))[:6000]
    except Exception as exc:  # pragma: no cover - depends on local docs
        result["error"] = repr(exc)
    return result


def extract_excel(path: Path) -> dict[str, Any]:
    result: dict[str, Any] = {"sheets": [], "sample_text": "", "error": None}
    try:
        import pandas as pd  # type: ignore

        workbook = pd.ExcelFile(path)
        pieces: list[str] = []
        for sheet in workbook.sheet_names:
            frame = pd.read_excel(path, sheet_name=sheet, nrows=20)
            result["sheets"].append(
                {
                    "name": sheet,
                    "rows_preview": len(frame),
                    "cols": [str(column) for column in frame.columns],
                }
            )
            pieces.append(f"[{sheet}]\n{frame.to_csv(index=False)}")
        result["sample_text"] = clean_text("\n".join(pieces))[:6000]
    except Exception as exc:  # pragma: no cover - depends on local sheets
        result["error"] = repr(exc)
    return result


def extract_rtf(path: Path) -> dict[str, Any]:
    result: dict[str, Any] = {"sample_text": "", "error": None}
    try:
        raw = path.read_text(encoding="utf-8", errors="ignore")
        # Light RTF cleanup; good enough for keyword indexing.
        raw = re.sub(r"\\'[0-9a-fA-F]{2}", " ", raw)
        raw = re.sub(r"\\[a-zA-Z]+-?\d* ?", " ", raw)
        raw = raw.replace("{", " ").replace("}", " ")
        result["sample_text"] = clean_text(raw)[:6000]
    except Exception as exc:
        result["error"] = repr(exc)
    return result


def extract_file(path: Path, max_pages: int) -> dict[str, Any]:
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        return extract_pdf(path, max_pages)
    if suffix == ".docx":
        return extract_docx(path)
    if suffix in {".xlsx", ".xls"}:
        return extract_excel(path)
    if suffix == ".rtf":
        return extract_rtf(path)
    return {"sample_text": "", "error": "unsupported extension"}


def build_index(source: Path, output_dir: Path, max_pages: int) -> dict[str, Any]:
    output_dir.mkdir(parents=True, exist_ok=True)
    records: list[dict[str, Any]] = []
    seen_hashes: dict[str, str] = {}

    for path in sorted(source.rglob("*")):
        if not path.is_file() or path.suffix.lower() not in SUPPORTED_EXTENSIONS:
            continue
        digest = file_hash(path)
        duplicate_of = seen_hashes.get(digest)
        if duplicate_of is None:
            seen_hashes[digest] = str(path)

        extracted = extract_file(path, max_pages=max_pages)
        sample_text = extracted.get("sample_text", "")
        records.append(
            {
                "path": str(path),
                "name": path.name,
                "extension": path.suffix.lower(),
                "size_bytes": path.stat().st_size,
                "modified": datetime.fromtimestamp(path.stat().st_mtime).isoformat(timespec="seconds"),
                "md5": digest,
                "duplicate_of": duplicate_of,
                "themes": infer_themes(path.name, sample_text),
                "sample_text": sample_text[:1200],
                "extract_meta": {key: value for key, value in extracted.items() if key != "sample_text"},
            }
        )

    unique = [record for record in records if not record["duplicate_of"]]
    theme_counts = Counter(theme for record in unique for theme in record["themes"])
    extension_counts = Counter(record["extension"] for record in records)
    index = {
        "root": str(source),
        "generated_at": datetime.now().isoformat(timespec="seconds"),
        "count": len(records),
        "unique_count": len(unique),
        "extension_counts": dict(extension_counts),
        "theme_counts": dict(theme_counts),
        "records": records,
    }

    (output_dir / "mason_library_index.json").write_text(
        json.dumps(index, ensure_ascii=False, indent=2), encoding="utf-8"
    )
    (output_dir / "mason_library_study_map.md").write_text(
        render_markdown(index), encoding="utf-8"
    )
    return index


def render_markdown(index: dict[str, Any]) -> str:
    records = index["records"]
    unique = [record for record in records if not record["duplicate_of"]]
    lines: list[str] = [
        "# 梅森资料库索引",
        "",
        f"- 根目录：`{index['root']}`",
        f"- 生成时间：{index['generated_at']}",
        f"- 文件数：{index['count']}，去重后：{index['unique_count']}",
        "",
        "## 主题分布",
        "",
    ]
    for theme, count in sorted(index["theme_counts"].items(), key=lambda item: item[1], reverse=True):
        lines.append(f"- {theme}: {count}")

    priority = [
        "宏观/策略/风格",
        "AI算力/国产算力",
        "CPO/光通信",
        "PCB/ABF/载板",
        "半导体/材料/设备",
        "小金属/有色/稀土",
        "能源/化工/油价",
        "机器人/无人驾驶",
        "AI应用/Agent",
    ]
    lines += ["", "## 建议优先学习路线", ""]
    for number, theme in enumerate(priority, 1):
        names = [record["name"] for record in unique if theme in record["themes"]][:8]
        if not names:
            continue
        lines.append(f"### {number}. {theme}")
        for name in names:
            lines.append(f"- {name}")

    duplicates = [record for record in records if record["duplicate_of"]]
    lines += ["", "## 去重提示", ""]
    if duplicates:
        for record in duplicates:
            lines.append(f"- `{record['name']}` 重复于 `{record['duplicate_of']}`")
    else:
        lines.append("- 暂无重复文件")

    lines += ["", "## 样本文档摘要预览", ""]
    for record in unique[:30]:
        preview = record["sample_text"][:220] or "未抽取到文本"
        lines.append(f"### {record['name']}")
        lines.append(f"- 主题：{', '.join(record['themes'])}")
        lines.append(f"- 预览：{preview}")
    return "\n".join(lines)


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="Build Mason local knowledge-library index.")
    parser.add_argument("--source", default=r"D:\梅森", help="Local Mason source folder.")
    parser.add_argument(
        "--output",
        default=str(Path("data") / "mason_library"),
        help="Output directory for JSON and Markdown index.",
    )
    parser.add_argument("--max-pages", type=int, default=3, help="PDF pages to sample per document.")
    return parser.parse_args()


def main() -> int:
    args = parse_args()
    source = Path(args.source)
    if not source.exists():
        raise SystemExit(f"Source folder not found: {source}")
    index = build_index(source=source, output_dir=Path(args.output), max_pages=args.max_pages)
    print(f"indexed={index['count']} unique={index['unique_count']}")
    print(Path(args.output) / "mason_library_index.json")
    print(Path(args.output) / "mason_library_study_map.md")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
